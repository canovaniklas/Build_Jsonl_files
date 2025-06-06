#!/usr/bin/env python3
"""
Convert a folder of alternating Romansh/German .docx files to a gzipped
JSON-Lines corpus.

• index 0 → Romansh ("rm")   • index 1 → German ("de")  … and so on
• Extracts **all** text: body paragraphs, tables, headers, footers, text-boxes.
"""

import argparse, gzip, json, uuid, pathlib, sys

def extract_with_docx2txt(path: pathlib.Path) -> str | None:
    """Try docx2txt (grabs headers/footers/tables/text-boxes)."""
    try:
        import docx2txt
        return docx2txt.process(str(path))
    except ImportError:
        return None
    except Exception:
        # any weird parsing error → fall back
        return None


def extract_with_python_docx(path: pathlib.Path) -> str:
    """
    Use python-docx and include *all* table cell texts in addition
    to regular paragraphs. Order of tables vs. paragraphs may differ
    from the original layout, but no text is lost.
    """
    from docx import Document

    doc = Document(path)
    chunks: list[str] = []

    # body paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            chunks.append(p.text)

    # tables (all cells)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if cell.text.strip():
                    chunks.append(cell.text)

    # headers / footers
    for sect in doc.sections:
        for hf in [sect.header, sect.footer]:
            for p in hf.paragraphs:
                if p.text.strip():
                    chunks.append(p.text)

    return "\n".join(chunks)


def docx_to_text(path: pathlib.Path) -> str:
    txt = extract_with_docx2txt(path)
    if txt:
        return txt
    # Guaranteed fallback:
    return extract_with_python_docx(path)


def lang_for(i: int) -> str:
    """rm, de, rm, de, … based on file index."""
    return "rm" if i % 2 == 0 else "de"


def build_record(path: pathlib.Path, idx: int) -> dict:
    return {
        "text": docx_to_text(path),
        "id": str(uuid.uuid4()),
        "metadata": {
            "filename": path.name,
            "language": lang_for(idx),
            "language_script": "Latn",
            "source": "Gemeinde Sagogn",
        },
    }

def main(src: str, out: str):
    src_root = pathlib.Path(src).expanduser().resolve()
    if not src_root.is_dir():
        sys.exit(f"Source directory not found: {src_root}")

    files = sorted(src_root.glob("*.docx"))
    if not files:
        sys.exit("No .docx files found in the source directory.")

    out_path = pathlib.Path(out).expanduser().resolve()
    out_path.parent.mkdir(parents=True, exist_ok=True)

    with gzip.open(out_path, "wt", encoding="utf-8") as gz:
        for idx, path in enumerate(files):
            rec = build_record(path, idx)
            json.dump(rec, gz, ensure_ascii=False)
            gz.write("\n")

    print(f"Wrote {out_path}  ({len(files)} documents)")

if __name__ == "__main__":
    ap = argparse.ArgumentParser(
        description="Bundle alternating rm/de .docx files into gzipped JSON-Lines."
    )
    ap.add_argument("--src", required=True, help="Folder that holds the .docx files")
    ap.add_argument("--out", required=True, help="Output file, e.g. corpus.jsonl.gz")
    args = ap.parse_args()
    main(src=args.src, out=args.out)
